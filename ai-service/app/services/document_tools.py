"""
文档知识管家 Agent 工具函数
提供文档解析、分段、摘要、知识结构、向量化等工具
"""

from typing import Dict, List, Any, Optional
from app.core.app_logging import logger
from app.core.llm_client import llm_client
from app.services.vector_store import vector_store


class DocumentTools:
    """文档处理工具类"""
    
    async def parse_document(
        self,
        file_path: str,
        file_type: str
    ) -> Dict[str, Any]:
        """
        解析文档工具
        
        解析 PDF/DOCX/TXT 等格式文档，提取文本内容
        
        Args:
            file_path: 文件路径
            file_type: 文件类型 (pdf, docx, txt, md)
            
        Returns:
            dict: {
                "raw_text": str,  # 原始文本
                "metadata": dict   # 文档元数据 (页数、标题等)
            }
        """
        logger.info(f"解析文档: file_path={file_path}, file_type={file_type}")
        
        try:
            # 检查文件是否存在
            import os
            if not os.path.exists(file_path):
                error_msg = f"文件不存在: {file_path}"
                logger.error(error_msg)
                raise FileNotFoundError(error_msg)
            
            # 根据文件类型选择对应的解析方法
            file_type_lower = file_type.lower().strip('.')
            
            if file_type_lower == "pdf":
                # 解析 PDF 文件
                return await self._parse_pdf(file_path)
            elif file_type_lower == "docx":
                # 解析 DOCX 文件
                return await self._parse_docx(file_path)
            elif file_type_lower in ["txt", "md"]:
                # 解析文本文件
                return await self._parse_text(file_path, file_type_lower)
            else:
                # 不支持的文件类型
                error_msg = f"不支持的文件类型: {file_type}"
                logger.error(error_msg)
                raise ValueError(error_msg)
                
        except FileNotFoundError:
            # 重新抛出文件不存在错误
            raise
        except Exception as e:
            logger.error(f"解析文档失败: {str(e)}", exc_info=True)
            raise
    
    async def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        解析 PDF 文件
        
        Args:
            file_path: PDF 文件路径
            
        Returns:
            dict: 包含原始文本和元数据
        """
        try:
            from pypdf import PdfReader
            import os
            
            logger.info(f"开始解析 PDF: {file_path}")
            
            # 使用 pypdf 读取 PDF
            reader = PdfReader(file_path)
            
            # 提取所有页面的文本
            pages_text = []
            total_pages = len(reader.pages)
            
            for page_num, page in enumerate(reader.pages, start=1):
                try:
                    text = page.extract_text()
                    if text.strip():  # 只添加非空页面
                        pages_text.append(text)
                        logger.debug(f"提取第 {page_num} 页文本，长度: {len(text)}")
                except Exception as e:
                    logger.warning(f"提取第 {page_num} 页文本失败: {str(e)}")
                    # 继续处理下一页
                    continue
            
            # 合并所有页面文本，用换行符分隔
            raw_text = "\n\n".join(pages_text)
            
            # 提取元数据
            metadata = reader.metadata if reader.metadata else {}
            title = metadata.get("/Title", "") if metadata else ""
            if not title:
                title = os.path.basename(file_path) or "未命名文档"
            
            # 构建返回的元数据
            result_metadata = {
                "file_type": "pdf",
                "file_path": file_path,
                "title": title,
                "total_pages": total_pages,
                "extracted_pages": len(pages_text),
                "author": metadata.get("/Author", "") if metadata else "",
                "subject": metadata.get("/Subject", "") if metadata else "",
            }
            
            logger.info(f"PDF 解析完成: text_length={len(raw_text)}, pages={total_pages}")
            
            return {
                "raw_text": raw_text,
                "metadata": result_metadata
            }
            
        except ImportError:
            error_msg = "pypdf 库未安装，请运行: pip install pypdf"
            logger.error(error_msg)
            raise ImportError(error_msg)
        except Exception as e:
            logger.error(f"解析 PDF 失败: {str(e)}", exc_info=True)
            raise
    
    async def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """
        解析 DOCX 文件
        
        Args:
            file_path: DOCX 文件路径
            
        Returns:
            dict: 包含原始文本和元数据
        """
        try:
            from docx import Document
            import os
            
            logger.info(f"开始解析 DOCX: {file_path}")
            
            # 使用 python-docx 读取 DOCX
            doc = Document(file_path)
            
            # 提取所有段落文本
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # 只添加非空段落
                    paragraphs.append(text)
            
            # 合并所有段落，用换行符分隔
            raw_text = "\n\n".join(paragraphs)
            
            # 提取表格文本（如果有）
            tables_text = []
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    if any(row_cells):  # 只添加非空行
                        table_rows.append(" | ".join(row_cells))
                if table_rows:
                    tables_text.append("\n".join(table_rows))
            
            # 如果有表格，添加到文本末尾
            if tables_text:
                raw_text += "\n\n[表格内容]\n\n" + "\n\n".join(tables_text)
            
            # 提取元数据
            core_props = doc.core_properties
            title = core_props.title or os.path.basename(file_path) or "未命名文档"
            
            # 构建返回的元数据
            result_metadata = {
                "file_type": "docx",
                "file_path": file_path,
                "title": title,
                "total_paragraphs": len(paragraphs),
                "total_tables": len(doc.tables),
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
            }
            
            logger.info(f"DOCX 解析完成: text_length={len(raw_text)}, paragraphs={len(paragraphs)}")
            
            return {
                "raw_text": raw_text,
                "metadata": result_metadata
            }
            
        except ImportError:
            error_msg = "python-docx 库未安装，请运行: pip install python-docx"
            logger.error(error_msg)
            raise ImportError(error_msg)
        except Exception as e:
            logger.error(f"解析 DOCX 失败: {str(e)}", exc_info=True)
            raise
    
    async def _parse_text(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """
        解析文本文件（TXT/MD）
        
        Args:
            file_path: 文本文件路径
            file_type: 文件类型 (txt, md)
            
        Returns:
            dict: 包含原始文本和元数据
        """
        import os
        
        logger.info(f"开始解析文本文件: {file_path}")
        
        try:
            # 尝试 UTF-8 编码
            with open(file_path, "r", encoding="utf-8") as f:
                raw_text = f.read()
        except UnicodeDecodeError:
            # 如果 UTF-8 失败，尝试 GBK 编码（常见于中文 Windows 系统）
            logger.warning(f"UTF-8 解码失败，尝试 GBK 编码: {file_path}")
            try:
                with open(file_path, "r", encoding="gbk") as f:
                    raw_text = f.read()
            except UnicodeDecodeError:
                # 如果 GBK 也失败，尝试 latin-1（几乎总能成功，但可能乱码）
                logger.warning(f"GBK 解码也失败，尝试 latin-1 编码: {file_path}")
                with open(file_path, "r", encoding="latin-1") as f:
                    raw_text = f.read()
        
        # 构建元数据
        result_metadata = {
            "file_type": file_type,
            "file_path": file_path,
            "title": os.path.basename(file_path) or "未命名文档",
        }
        
        logger.info(f"文本文件解析完成: text_length={len(raw_text)}")
        
        return {
            "raw_text": raw_text,
            "metadata": result_metadata
        }
    
    async def chunk_text(
        self,
        text: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200
    ) -> Dict[str, Any]:
        """
        文本分段工具
        
        将长文本按照固定大小分段，支持重叠
        
        Args:
            text: 文本内容
            chunk_size: 分段大小（字符数）
            chunk_overlap: 重叠大小（字符数）
            
        Returns:
            dict: {
                "chunks": List[str],  # 分段列表
                "chunk_metadata": List[dict]  # 每个分段的元数据
            }
        """
        logger.info(f"文本分段: text_length={len(text)}, chunk_size={chunk_size}, overlap={chunk_overlap}")
        
        try:
            # 如果文本为空或很短，直接返回单个分段
            if not text or len(text) <= chunk_size:
                return {
                    "chunks": [text] if text else [],
                    "chunk_metadata": [{
                        "index": 0,
                        "start": 0,
                        "end": len(text),
                        "length": len(text)
                    }] if text else []
                }
            
            chunks = []
            chunk_metadata = []
            
            # 确保 chunk_overlap 不会导致无限循环
            # 如果 overlap 太大，限制为 chunk_size 的一半
            if chunk_overlap >= chunk_size:
                chunk_overlap = max(1, chunk_size // 2)
                logger.warning(f"chunk_overlap 过大，已调整为: {chunk_overlap}")
            
            # 简单实现：按固定大小分段
            start = 0
            index = 0
            
            while start < len(text):
                end = min(start + chunk_size, len(text))
                chunk = text[start:end]
                chunks.append(chunk)
                
                chunk_metadata.append({
                    "index": index,
                    "start": start,
                    "end": end,
                    "length": len(chunk)
                })
                
                # 移动到下一个分段（考虑重叠）
                # 确保 start 总是递增，避免无限循环
                next_start = end - chunk_overlap
                if next_start <= start:
                    # 如果下一个 start 没有递增，强制递增至少 1
                    next_start = start + 1
                
                # 如果已经到达文本末尾，退出循环
                if end >= len(text):
                    break
                    
                start = next_start
                index += 1
                
                # 安全保护：如果分段数量过多，停止处理
                if index > 10000:
                    logger.error(f"分段数量过多 ({index})，可能存在无限循环，停止处理")
                    raise ValueError(f"分段数量过多，可能存在逻辑错误")
            
            logger.info(f"文本分段完成: chunks_count={len(chunks)}")
            
            return {
                "chunks": chunks,
                "chunk_metadata": chunk_metadata
            }
        except Exception as e:
            logger.error(f"文本分段失败: {str(e)}", exc_info=True)
            raise
    
    async def generate_summary(
        self,
        text: str,
        summary_type: str = "document"
    ) -> Dict[str, Any]:
        """
        生成摘要工具
        
        使用 LLM 生成文档或分段的摘要
        
        Args:
            text: 文本内容
            summary_type: 摘要类型 ("document" 或 "chunk")
            
        Returns:
            dict: {
                "summary": str,  # 摘要文本
                "key_points": List[str]  # 关键点列表（可选）
            }
        """
        logger.info(f"生成摘要: text_length={len(text)}, summary_type={summary_type}")
        
        try:
            # 构建提示词
            if summary_type == "chunk":
                system_prompt = "你是一个专业的文本摘要助手。请为给定的文本片段生成简洁准确的摘要，控制在100字以内。"
                user_prompt = f"请为以下文本片段生成摘要：\n\n{text}"
            else:
                system_prompt = "你是一个专业的文档摘要助手。请为给定的文档生成全面准确的摘要，控制在300字以内，并提取3-5个关键点。"
                user_prompt = f"请为以下文档生成摘要和关键点：\n\n{text}"
            
            # 调用 LLM
            result = await llm_client.chat(
                prompt=user_prompt,
                system_prompt=system_prompt,
                temperature=0.3,  # 降低温度以获得更准确的摘要
                max_tokens=500 if summary_type == "document" else 200
            )
            
            summary = result["reply"].strip()
            
            # 尝试提取关键点（从摘要中提取或单独生成）
            key_points = []
            if summary_type == "document":
                # 尝试从回复中提取关键点（如果 LLM 返回了列表格式）
                lines = summary.split('\n')
                for line in lines:
                    line = line.strip()
                    # 检查是否是关键点格式（以数字、-、• 开头）
                    if line and (line[0].isdigit() or line.startswith('-') or line.startswith('•')):
                        key_point = line.lstrip('0123456789.-• ').strip()
                        if key_point:
                            key_points.append(key_point)
            
            # 如果没有提取到关键点，且是文档摘要，尝试单独生成
            if summary_type == "document" and not key_points:
                try:
                    key_points_prompt = f"请从以下文档中提取3-5个关键点，每行一个：\n\n{text[:2000]}"  # 限制长度
                    key_points_result = await llm_client.chat(
                        prompt=key_points_prompt,
                        system_prompt="你是一个专业的信息提取助手。请从文档中提取关键点，每行一个。",
                        temperature=0.3,
                        max_tokens=200
                    )
                    key_points = [
                        line.strip().lstrip('0123456789.-• ')
                        for line in key_points_result["reply"].split('\n')
                        if line.strip() and not line.strip().startswith('#')
                    ][:5]  # 最多5个
                except Exception as e:
                    logger.warning(f"提取关键点失败: {str(e)}")
            
            logger.info(f"摘要生成完成: summary_length={len(summary)}, key_points_count={len(key_points)}")
            
            return {
                "summary": summary,
                "key_points": key_points
            }
        except Exception as e:
            logger.error(f"生成摘要失败: {str(e)}", exc_info=True)
            # 失败时返回简单截断作为后备
            max_length = 200 if summary_type == "chunk" else 500
            summary = text[:max_length] + "..." if len(text) > max_length else text
            return {
                "summary": summary,
                "key_points": []
            }
    
    async def build_knowledge_structure(
        self,
        chunks: List[str],
        document_title: str
    ) -> str:
        """
        构建知识结构工具
        
        分析文档内容，构建知识图谱或层次结构
        
        Args:
            chunks: 分段列表
            document_title: 文档标题
            
        Returns:
            str: 知识结构 (JSON 字符串)
        """
        logger.info(f"构建知识结构: document_title={document_title}, chunks_count={len(chunks)}")
        
        try:
            # TODO: 实现真实的知识结构构建逻辑
            # 1. 分析文档内容，提取主题、章节、实体等
            # 2. 构建层次结构或知识图谱
            # 3. 返回 JSON 格式的知识结构
            
            # 当前为简单实现：返回基本结构
            knowledge_structure = {
                "title": document_title,
                "sections": [
                    {
                        "title": f"章节 {i+1}",
                        "chunk_indices": [i]
                    }
                    for i in range(len(chunks))
                ],
                "topics": [],
                "relationships": []
            }
            
            import json
            structure_json = json.dumps(knowledge_structure, ensure_ascii=False)
            
            logger.info("知识结构构建完成")
            
            return structure_json
        except Exception as e:
            logger.error(f"构建知识结构失败: {str(e)}", exc_info=True)
            raise
    
    async def embed_text(
        self,
        text: str,
        model_name: str = "text-embedding-ada-002"
    ) -> Dict[str, Any]:
        """
        向量化工具
        
        将文本转换为向量表示
        
        Args:
            text: 文本内容
            model_name: 模型名称
            
        Returns:
            dict: {
                "vector": List[float],  # 向量数据
                "dimension": int  # 向量维度
            }
        """
        logger.info(f"向量化文本: text_length={len(text)}, model={model_name}")
        
        try:
            # TODO: 实现真实的向量化逻辑
            # 1. 调用 Embedding API（如 OpenAI Embeddings）
            # 2. 返回向量数据
            
            # 当前为简单实现：返回随机向量（仅用于测试）
            dimension = 1536  # OpenAI ada-002 的维度
            import random
            vector = [random.random() for _ in range(dimension)]
            
            logger.info(f"向量化完成: dimension={dimension}")
            
            return {
                "vector": vector,
                "dimension": dimension
            }
        except Exception as e:
            logger.error(f"向量化失败: {str(e)}", exc_info=True)
            raise
    
    async def store_vectors(
        self,
        vectors: List[List[float]],
        metadata: List[Dict[str, Any]]
    ) -> Dict[str, Any]:
        """
        向量存储工具
        
        将向量存储到向量数据库
        
        Args:
            vectors: 向量列表
            metadata: 元数据列表
            
        Returns:
            dict: {
                "success": bool,
                "stored_count": int,
                "vector_ids": List[str]  # 存储后的向量 ID 列表
            }
        """
        logger.info(f"存储向量: vectors_count={len(vectors)}")
        
        try:
            # 从 metadata 中提取 document_id（用于组织集合）
            document_id = None
            if metadata and "document_id" in metadata[0]:
                document_id = str(metadata[0]["document_id"])
            
            # 准备文本列表（用于存储到 Chroma）
            # 如果 metadata 中有 content，使用它；否则使用 summary 或空字符串
            texts = []
            for meta in metadata:
                text = meta.get("content", meta.get("summary", ""))
                texts.append(text)
            
            # 调用 vector_store 存储向量
            vector_ids = await vector_store.add_vectors(
                vectors=vectors,
                texts=texts,
                metadata=metadata,
                document_id=document_id
            )
            
            logger.info(f"向量存储完成: stored_count={len(vector_ids)}, document_id={document_id}")
            
            return {
                "success": True,
                "stored_count": len(vector_ids),
                "vector_ids": vector_ids
            }
        except Exception as e:
            logger.error(f"向量存储失败: {str(e)}", exc_info=True)
            raise
    
    async def search_vectors(
        self,
        query_vector: List[float],
        top_k: int = 5,
        document_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        向量检索工具
        
        在向量数据库中检索相似文档片段
        
        Args:
            query_vector: 查询向量
            top_k: 返回数量
            document_id: 文档 ID（可选，限定文档）
            
        Returns:
            dict: {
                "results": List[dict]  # 检索结果列表
                    {
                        "chunk_id": int,
                        "document_id": str,
                        "score": float,
                        "content": str,
                        "summary": str,
                        "chunk_index": int
                    }
            }
        """
        logger.info(f"向量检索: top_k={top_k}, document_id={document_id}")
        
        try:
            # 调用 vector_store 检索向量
            search_results = await vector_store.search_vectors(
                query_vector=query_vector,
                top_k=top_k,
                document_id=document_id
            )
            
            # 转换结果格式，匹配期望的输出格式
            results = []
            for item in search_results:
                metadata = item.get("metadata", {})
                
                # 确保 score 在 0-1 范围内
                raw_score = item.get("score", 0.0)
                
                # 调试日志：如果 score 异常，记录详细信息
                if raw_score < 0.0 or raw_score > 1.0:
                    logger.warning(f"⚠️ document_tools.search_vectors: 检测到异常的 score 值: {raw_score}, item keys: {item.keys()}")
                    logger.warning(f"   原始 item: {item}")
                
                # 强制修正 score 到 0-1 范围
                if raw_score < 0.0:
                    score = 0.0
                    logger.warning(f"   Score 为负数 {raw_score}，已修正为 0.0")
                elif raw_score > 1.0:
                    score = 1.0
                    logger.warning(f"   Score 超过 1.0 ({raw_score})，已修正为 1.0")
                else:
                    score = raw_score
                
                # 再次确保 score 在有效范围内（双重保险）
                final_score = max(0.0, min(1.0, float(score)))
                
                result_item = {
                    "chunk_id": int(metadata.get("chunk_id", metadata.get("chunk_index", 0))),
                    "document_id": str(metadata.get("document_id", document_id or "unknown")),
                    "score": final_score,
                    "content": item.get("document", metadata.get("content", "")),
                    "summary": metadata.get("summary"),
                    "chunk_index": int(metadata.get("chunk_index", metadata.get("chunk_id", 0)))
                }
                results.append(result_item)
            
            logger.info(f"向量检索完成: results_count={len(results)}")
            
            return {
                "results": results
            }
        except Exception as e:
            logger.error(f"向量检索失败: {str(e)}", exc_info=True)
            raise
    
    async def generate_answer(
        self,
        query: str,
        context: str,
        document_id: str
    ) -> Dict[str, Any]:
        """
        问答生成工具
        
        基于检索到的文档片段，使用 LLM 生成答案
        
        Args:
            query: 用户问题
            context: 上下文文档片段
            document_id: 文档 ID
            
        Returns:
            dict: {
                "answer": str,  # 答案文本
                "confidence": float  # 置信度 (0-1)
            }
        """
        logger.info(f"生成答案: query={query[:50]}..., context_length={len(context)}")
        
        try:
            # 构建提示词
            system_prompt = """你是一个专业的文档问答助手。请根据提供的文档片段内容，准确回答用户的问题。

要求：
1. 答案必须基于提供的文档内容，不要编造信息
2. 如果文档中没有相关信息，请明确说明"根据提供的文档内容，无法找到相关信息"
3. 答案要简洁明了，重点突出
4. 如果涉及多个片段的内容，请综合整理"""
            
            user_prompt = f"""用户问题：{query}

相关文档片段：
{context}

请根据以上文档片段回答用户的问题。"""
            
            # 调用 LLM
            result = await llm_client.chat(
                prompt=user_prompt,
                system_prompt=system_prompt,
                temperature=0.3,  # 降低温度以获得更准确的答案
                max_tokens=1000
            )
            
            answer = result["reply"].strip()
            
            # 计算置信度（简单启发式方法）
            # 如果答案中包含"无法找到"、"没有相关信息"等，置信度较低
            low_confidence_keywords = ["无法找到", "没有相关信息", "文档中没有", "无法确定"]
            confidence = 0.3 if any(keyword in answer for keyword in low_confidence_keywords) else 0.85
            
            # 如果答案太短，可能置信度较低
            if len(answer) < 20:
                confidence = max(0.3, confidence - 0.2)
            
            logger.info(f"答案生成完成: answer_length={len(answer)}, confidence={confidence}")
            
            return {
                "answer": answer,
                "confidence": confidence
            }
        except Exception as e:
            logger.error(f"生成答案失败: {str(e)}", exc_info=True)
            # 失败时返回默认答案
            return {
                "answer": f"抱歉，生成答案时出现错误：{str(e)}。请稍后重试。",
                "confidence": 0.1
            }

